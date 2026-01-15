"""
Docx MCP Server using FastMCP and python-docx.

This server allows ChatGPT to read, search, and modify DOCX files.
Supports:
- Dynamic document switching with fuzzy search
- Reading and searching paragraphs
- Inserting content at arbitrary positions with formatting (bold, italic, alignment, styles)
- Template placeholders (<<Name>>, {{Date}}, etc.)
- Find and replace text for surgical edits
- Automatic table detection and conversion (markdown & tab-delimited)
- Full table manipulation (read, update cells, add/delete rows)
"""

import os
import re
import copy
from difflib import SequenceMatcher
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips, Cm
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from fastmcp import FastMCP

# Define the default document path (now in documents/ folder)
DEFAULT_DOCX_PATH = os.environ.get("DOCX_PATH", os.path.join("documents", "MCP.docx"))

# Global variable to track current active document
CURRENT_DOCX_PATH = DEFAULT_DOCX_PATH


# Create the MCP server
mcp = FastMCP(name="Docx Editor")


def get_document():
    """Helper to load the document or create a new one if it doesn't exist."""
    global CURRENT_DOCX_PATH
    if os.path.exists(CURRENT_DOCX_PATH):
        return Document(CURRENT_DOCX_PATH)
    return Document()


def save_document(doc):
    """Save document to the current active path."""
    global CURRENT_DOCX_PATH
    doc.save(CURRENT_DOCX_PATH)


def find_document_by_name(query: str, search_dir: str = ".") -> list:
    """
    Find .docx files by fuzzy name matching in the specified directory.
    Returns list of (file_path, filename, score) tuples sorted by score.
    """
    results = []

    # Walk through directory to find all .docx files
    for root, dirs, files in os.walk(search_dir):
        for file in files:
            if file.endswith('.docx') and not file.startswith('~$'):  # Skip temp files
                file_path = os.path.join(root, file)
                filename = os.path.splitext(file)[0]  # Remove .docx extension

                # Calculate similarity
                score = similarity(query.lower(), filename.lower())

                # Also check if query is contained in filename
                if query.lower() in filename.lower():
                    score = max(score, 0.8)

                if score >= 0.3:  # Minimum threshold
                    results.append((file_path, file, score))

    # Sort by score descending
    results.sort(key=lambda x: x[2], reverse=True)
    return results


def similarity(a: str, b: str) -> float:
    """Calculate similarity ratio between two strings using SequenceMatcher."""
    return SequenceMatcher(None, a.lower(), b.lower()).ratio()


def find_paragraph_by_text(doc, query: str, threshold: float = 0.5):
    """
    Find a paragraph by fuzzy text matching.
    Returns (index, paragraph, score) or None if not found.
    """
    best_match = None
    best_score = 0

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Check if query is contained in paragraph
        if query.lower() in text.lower():
            score = 0.9 + (0.1 * similarity(query, text))
            if score > best_score:
                best_score = score
                best_match = (idx, para, score)
        else:
            # Use fuzzy matching
            score = similarity(query, text)
            if score > best_score and score >= threshold:
                best_score = score
                best_match = (idx, para, score)

    return best_match


def find_placeholders(doc) -> list:
    """Find all placeholders in the document (<<...>> or {{...>>), including inside tables."""
    placeholders = []
    pattern = re.compile(r'(<<[^<>]+>>|\{\{[^{}]+\}\})')

    # Search in paragraphs
    for idx, para in enumerate(doc.paragraphs):
        matches = pattern.findall(para.text)
        for match in matches:
            placeholders.append({
                "placeholder": match,
                "location_type": "paragraph",
                "paragraph_index": idx,
                "context": para.text[:100] + "..." if len(para.text) > 100 else para.text
            })

    # Search in tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text
                matches = pattern.findall(cell_text)
                for match in matches:
                    placeholders.append({
                        "placeholder": match,
                        "location_type": "table",
                        "table_index": table_idx,
                        "row": row_idx,
                        "column": col_idx,
                        "context": cell_text[:100] + "..." if len(cell_text) > 100 else cell_text
                    })

    return placeholders


def replace_text_in_paragraph(paragraph, old_text: str, new_text: str) -> bool:
    """
    Replace text in a paragraph, handling the case where text might be split across runs.
    Returns True if replacement was made.
    """
    # First try simple replacement in the full text
    if old_text in paragraph.text:
        # Try to find and replace in individual runs first
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                return True

        # If not found in individual runs, the text is split across runs
        # We need to rebuild the paragraph
        full_text = paragraph.text
        new_full_text = full_text.replace(old_text, new_text)

        # Clear all runs and add the new text
        for run in paragraph.runs:
            run.text = ""

        if paragraph.runs:
            paragraph.runs[0].text = new_full_text
        else:
            paragraph.add_run(new_full_text)

        return True

    return False


def detect_table_format(text: str) -> tuple[str, list[list[str]]]:
    """
    Detect if text is a table and parse it.
    Returns (format_type, table_data) where format_type is 'markdown', 'tab', or None.
    table_data is a list of rows, where each row is a list of cell values.
    """
    lines = [line.strip() for line in text.strip().split('\n') if line.strip()]

    if not lines:
        return None, []

    # Check for markdown table format (| col1 | col2 |)
    if lines[0].startswith('|') and lines[0].endswith('|'):
        table_data = []
        for line in lines:
            # Skip separator lines like |------|------| or |:-----|-----:|
            # Check if line contains only dashes, colons, spaces, and pipes
            if re.match(r'^\s*\|[\s\-:|]+\|\s*$', line):
                # Further check: make sure it's mostly dashes (not actual content)
                content = line.replace('|', '').replace('-', '').replace(':', '').replace(' ', '')
                if len(content) == 0:
                    continue

            # Parse cells
            cells = [cell.strip() for cell in line.split('|')[1:-1]]

            # Skip if all cells are just dashes (another way separator lines appear)
            if all(re.match(r'^[\s\-:]+$', cell) for cell in cells if cell):
                continue

            if cells:
                table_data.append(cells)

        if len(table_data) >= 1:  # At least header (changed from 2 to 1)
            return 'markdown', table_data

    # Check for tab-delimited format
    tab_rows = []
    for line in lines:
        if '\t' in line:
            cells = [cell.strip() for cell in line.split('\t')]
            tab_rows.append(cells)

    if tab_rows and len(tab_rows) >= 2:
        # Check if all rows have similar column count
        col_counts = [len(row) for row in tab_rows]
        if max(col_counts) - min(col_counts) <= 1:  # Allow 1 column variance
            return 'tab', tab_rows

    return None, []


def apply_paragraph_formatting(paragraph, bold: bool = False, italic: bool = False,
                               alignment: str = None, style: str = None):
    """
    Apply formatting to a paragraph.

    Args:
        paragraph: The paragraph object to format
        bold: Make text bold
        italic: Make text italic
        alignment: Text alignment - "left", "center", "right", "justify"
        style: Word style name - "Heading 1", "Title", etc.
    """
    # Apply text formatting to all runs
    if bold or italic:
        for run in paragraph.runs:
            if bold:
                run.bold = True
            if italic:
                run.italic = True

    # Apply alignment
    if alignment:
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        if alignment.lower() in alignment_map:
            paragraph.alignment = alignment_map[alignment.lower()]

    # Apply style
    if style:
        try:
            paragraph.style = style
        except KeyError:
            # Style doesn't exist, ignore
            pass


def create_word_table(doc, table_data: list[list[str]], has_header: bool = True):
    """
    Create a Word table from parsed table data.
    Returns the created table object.
    """
    if not table_data or not table_data[0]:
        return None

    rows = len(table_data)
    cols = max(len(row) for row in table_data)

    # Create table
    table = doc.add_table(rows=rows, cols=cols)

    # Apply the most common table style, fall back to default if not available
    try:
        table.style = 'Table Grid'
    except KeyError:
        # If style not available, leave as default (no borders)
        pass

    # Fill in data
    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_value in enumerate(row_data):
            if j < len(row.cells):
                row.cells[j].text = cell_value

                # Make header row bold
                if i == 0 and has_header:
                    for paragraph in row.cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    return table


# ============================================
# Document Management Tools
# ============================================

@mcp.tool()
async def get_current_document() -> dict:
    """
    Get the name and path of the currently active document.
    """
    global CURRENT_DOCX_PATH
    return {
        "current_document": os.path.basename(CURRENT_DOCX_PATH),
        "full_path": os.path.abspath(CURRENT_DOCX_PATH),
        "exists": os.path.exists(CURRENT_DOCX_PATH)
    }


@mcp.tool()
async def list_documents(search_dir: str = ".") -> dict:
    """
    List all .docx files in the specified directory and subdirectories.

    Args:
        search_dir: Directory to search in (default is current directory)
    """
    documents = []

    for root, dirs, files in os.walk(search_dir):
        for file in files:
            if file.endswith('.docx') and not file.startswith('~$'):
                file_path = os.path.join(root, file)
                rel_path = os.path.relpath(file_path, search_dir)
                documents.append({
                    "filename": file,
                    "path": rel_path,
                    "full_path": os.path.abspath(file_path)
                })

    return {
        "total": len(documents),
        "documents": documents[:50]  # Limit to 50 results
    }


@mcp.tool()
async def switch_document(query: str, search_dir: str = ".") -> dict:
    """
    SWITCH to a different document by name using fuzzy search.
    This changes which document all other tools will operate on.

    Args:
        query: The document name to search for (fuzzy matching)
        search_dir: Directory to search in (default is current directory)
    """
    global CURRENT_DOCX_PATH

    # Find matching documents
    results = find_document_by_name(query, search_dir)

    if not results:
        return {
            "error": f"No .docx files found matching '{query}' in '{search_dir}'",
            "suggestion": "Use list_documents to see available files"
        }

    # Use the best match
    best_match = results[0]
    file_path, filename, score = best_match

    # Update the current document path
    CURRENT_DOCX_PATH = file_path

    # Return results with top matches
    top_matches = [
        {"filename": fn, "score": round(sc, 2), "path": fp}
        for fp, fn, sc in results[:5]
    ]

    return {
        "status": "success",
        "message": f"Switched to document: {filename}",
        "current_document": filename,
        "full_path": os.path.abspath(file_path),
        "match_score": round(score, 2),
        "other_matches": top_matches[1:] if len(top_matches) > 1 else []
    }


# ============================================
# ChatGPT Required Tools: search and fetch
# ============================================

@mcp.tool()
async def search(query: str) -> dict:
    """
    Search for paragraphs in the DOCX document using fuzzy text matching.
    Returns a list of matching paragraphs with IDs for retrieval.

    Args:
        query: The search query string
    """
    doc = get_document()
    results = []

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Check containment first
        if query.lower() in text.lower():
            score = 0.9
        else:
            score = similarity(query, text)

        if score >= 0.3:  # Lower threshold for more results
            results.append({
                "id": f"para-{idx}",
                "score": round(score, 2),
                "text": text[:200] + "..." if len(text) > 200 else text
            })

    # Sort by score descending
    results.sort(key=lambda x: x["score"], reverse=True)

    return {"results": results[:10]}


@mcp.tool()
async def fetch(id: str) -> dict:
    """
    Retrieve the full content of a paragraph by its ID.
    Use IDs returned from the search tool.

    Args:
        id: The paragraph ID (e.g., "para-5")
    """
    doc = get_document()

    # Parse the index from the ID
    if not id.startswith("para-"):
        return {"error": f"Invalid ID format: {id}. Expected 'para-N'"}

    try:
        idx = int(id.replace("para-", ""))
    except ValueError:
        return {"error": f"Invalid ID format: {id}"}

    if idx < 0 or idx >= len(doc.paragraphs):
        return {"error": f"Paragraph {id} not found"}

    para = doc.paragraphs[idx]

    return {
        "id": id,
        "index": idx,
        "text": para.text,
        "style": para.style.name if para.style else None
    }


# ============================================
# Reading Tools
# ============================================

@mcp.tool()
async def read_document() -> dict:
    """
    Read the full text content of the document.
    """
    doc = get_document()
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    return {"text": "\n".join(full_text)}


@mcp.tool()
async def get_paragraphs(limit: int = 50, start_index: int = 0) -> dict:
    """
    Get a list of paragraphs from the document with their IDs.

    Args:
        limit: Maximum number of paragraphs to return (default 50)
        start_index: Index to start reading from (default 0)
    """
    doc = get_document()
    paragraphs = []

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append({
                "id": f"para-{idx}",
                "text": text[:100] + "..." if len(text) > 100 else text
            })

    # Slice the result
    sliced = paragraphs[start_index : start_index + limit]

    return {
        "total_paragraphs": len(paragraphs),
        "start_index": start_index,
        "showing": len(sliced),
        "paragraphs": sliced
    }


# ============================================
# Editing Tools
# ============================================

@mcp.tool()
async def add_paragraph(
    text: str,
    bold: bool = False,
    italic: bool = False,
    alignment: str = None,
    style: str = None
) -> dict:
    """
    Append a new paragraph to the end of the document with optional formatting.

    Args:
        text: Text content to add
        bold: Make text bold (default False)
        italic: Make text italic (default False)
        alignment: Text alignment - "left", "center", "right", "justify" (default left)
        style: Word style name - "Heading 1", "Heading 2", "Title", "Subtitle", etc.
    """
    doc = get_document()
    para = doc.add_paragraph(text)

    # Apply formatting
    apply_paragraph_formatting(para, bold=bold, italic=italic, alignment=alignment, style=style)

    save_document(doc)
    return {"status": "success", "message": "Paragraph added to end of document."}


@mcp.tool()
async def update_paragraph(
    id: str,
    text: str,
    bold: bool = None,
    italic: bool = None,
    alignment: str = None,
    style: str = None
) -> dict:
    """
    Update the text content of a paragraph by its ID with optional formatting.
    This replaces the entire paragraph text.

    Args:
        id: The paragraph ID to update (e.g., "para-5")
        text: The new text content for the paragraph
        bold: Make text bold (None = don't change)
        italic: Make text italic (None = don't change)
        alignment: Text alignment - "left", "center", "right", "justify" (None = don't change)
        style: Word style name - "Heading 1", "Title", etc. (None = don't change)
    """
    doc = get_document()

    # Parse the index from the ID
    if not id.startswith("para-"):
        return {"error": f"Invalid ID format: {id}. Expected 'para-N'"}

    try:
        idx = int(id.replace("para-", ""))
    except ValueError:
        return {"error": f"Invalid ID format: {id}"}

    if idx < 0 or idx >= len(doc.paragraphs):
        return {"error": f"Paragraph {id} not found"}

    para = doc.paragraphs[idx]

    # Clear existing runs and set new text
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)

    # Apply formatting if specified (using False instead of None to actually apply)
    if bold is not None or italic is not None or alignment or style:
        apply_paragraph_formatting(
            para,
            bold=bold if bold is not None else False,
            italic=italic if italic is not None else False,
            alignment=alignment,
            style=style
        )

    save_document(doc)
    return {"status": "success", "message": f"Paragraph {id} updated."}


@mcp.tool()
async def insert_before_text(
    query: str,
    text: str,
    threshold: float = 0.5,
    bold: bool = False,
    italic: bool = False,
    alignment: str = None,
    style: str = None
) -> dict:
    """
    INSERT CONTENT BEFORE a specific paragraph in the document with optional formatting.
    Finds any paragraph by text search and inserts new content immediately before it.

    Use this tool whenever you need to:
    - Insert before a specific section
    - Add content before certain text
    - Place new paragraphs at arbitrary positions (not just at the end)

    Args:
        query: The text to search for - finds the paragraph containing or matching this text
        text: The text content to insert as a new paragraph before the found location
        threshold: Minimum similarity score 0-1 (default 0.5)
        bold: Make text bold (default False)
        italic: Make text italic (default False)
        alignment: Text alignment - "left", "center", "right", "justify" (default left)
        style: Word style name - "Heading 1", "Title", etc.
    """
    doc = get_document()

    match = find_paragraph_by_text(doc, query, threshold)

    if not match:
        return {"error": f"No paragraph found matching: '{query}'"}

    idx, para, score = match

    # Insert a new paragraph before the found one
    new_para = doc.add_paragraph(text)

    # Move the new paragraph to the correct position (before the found paragraph)
    para._element.addprevious(new_para._element)

    # Apply formatting
    apply_paragraph_formatting(new_para, bold=bold, italic=italic, alignment=alignment, style=style)

    save_document(doc)

    return {
        "status": "success",
        "message": f"Content inserted before paragraph {idx} (match score: {score:.2f})",
        "matched_text": para.text[:100] + "..." if len(para.text) > 100 else para.text
    }


@mcp.tool()
async def insert_after_text(
    query: str,
    text: str,
    threshold: float = 0.5,
    bold: bool = False,
    italic: bool = False,
    alignment: str = None,
    style: str = None
) -> dict:
    """
    INSERT CONTENT AT ANY POSITION in the document with optional formatting.
    This is the PRIMARY tool for adding content at a specific location.
    Finds any paragraph by text search and inserts new content immediately after it.

    Use this tool whenever you need to:
    - Insert after a specific section
    - Add content after certain text
    - Place new paragraphs at arbitrary positions (not just at the end)

    Args:
        query: The text to search for - finds the paragraph containing or matching this text
        text: The text content to insert as a new paragraph after the found location
        threshold: Minimum similarity score 0-1 (default 0.5)
        bold: Make text bold (default False)
        italic: Make text italic (default False)
        alignment: Text alignment - "left", "center", "right", "justify" (default left)
        style: Word style name - "Heading 1", "Title", etc.
    """
    doc = get_document()

    match = find_paragraph_by_text(doc, query, threshold)

    if not match:
        return {"error": f"No paragraph found matching: '{query}'"}

    idx, para, score = match

    # Insert a new paragraph after the found one
    # python-docx doesn't have a direct insert_after, so we need to work with the XML
    new_para = doc.add_paragraph(text)

    # Move the new paragraph to the correct position
    para._element.addnext(new_para._element)

    # Apply formatting
    apply_paragraph_formatting(new_para, bold=bold, italic=italic, alignment=alignment, style=style)

    save_document(doc)

    return {
        "status": "success",
        "message": f"Content inserted after paragraph {idx} (match score: {score:.2f})",
        "matched_text": para.text[:100] + "..." if len(para.text) > 100 else para.text
    }


@mcp.tool()
async def insert_after_heading(heading_text: str, text: str) -> dict:
    """
    INSERT CONTENT AFTER A SECTION/HEADING.
    Finds a heading by text and inserts new content at the end of that section
    (before the next heading of same or higher level).

    Args:
        heading_text: The heading/section title to find
        text: The text content to insert after that section
    """
    doc = get_document()

    # Find the heading
    heading_idx = None
    heading_level = None

    for idx, para in enumerate(doc.paragraphs):
        if para.style and para.style.name.startswith('Heading'):
            if heading_text.lower() in para.text.lower():
                heading_idx = idx
                # Extract heading level (e.g., "Heading 1" -> 1)
                try:
                    heading_level = int(para.style.name.replace('Heading ', ''))
                except ValueError:
                    heading_level = 1
                break

    if heading_idx is None:
        # Try fuzzy matching on all paragraphs if no heading style found
        match = find_paragraph_by_text(doc, heading_text, 0.6)
        if match:
            heading_idx = match[0]
            heading_level = None

    if heading_idx is None:
        return {"error": f"No heading found matching: '{heading_text}'"}

    # Find the end of this section (next heading of same or higher level)
    insert_idx = heading_idx

    if heading_level:
        for idx in range(heading_idx + 1, len(doc.paragraphs)):
            para = doc.paragraphs[idx]
            if para.style and para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.replace('Heading ', ''))
                    if level <= heading_level:
                        insert_idx = idx - 1
                        break
                except ValueError:
                    pass
            insert_idx = idx

    # Insert the new paragraph
    target_para = doc.paragraphs[insert_idx]
    new_para = doc.add_paragraph(text)
    target_para._element.addnext(new_para._element)

    save_document(doc)

    return {
        "status": "success",
        "message": f"Content inserted after section '{heading_text}' (at paragraph {insert_idx})"
    }


# ============================================
# Template/Placeholder Tools
# ============================================

@mcp.tool()
async def list_placeholders() -> dict:
    """
    TEMPLATE SUPPORT - CALL THIS FIRST when working with templates.
    Find all placeholders in the document like <<Name>>, <<Poem>>, {{Date}}, etc.

    IMPORTANT: If a user mentions any placeholder (text with << >> or {{ }}),
    ALWAYS call this tool first to discover placeholders, then use replace_placeholder to fill them in.
    """
    doc = get_document()
    placeholders = find_placeholders(doc)

    # Get unique placeholder names
    unique = list(set(p["placeholder"] for p in placeholders))

    if not unique:
        return {
            "found": 0,
            "message": "No placeholders found. Placeholders should be formatted as <<Name>> or {{Name}}."
        }

    return {
        "found": len(unique),
        "placeholders": unique,
        "details": placeholders
    }


@mcp.tool()
async def replace_placeholder(placeholder: str, value: str) -> dict:
    """
    TEMPLATE SUPPORT - USE THIS TO FILL PLACEHOLDERS.
    Replaces a placeholder like <<Poem>> or {{Name}} with actual content IN-PLACE.
    This REMOVES the placeholder and puts the new content exactly where it was.
    Works in both paragraphs AND table cells.

    ALWAYS use this tool (not insert_after_text) when working with template placeholders.

    Args:
        placeholder: The EXACT placeholder to replace, including brackets (e.g., "<<Poem>>")
        value: The content to put in place of the placeholder
    """
    doc = get_document()
    count = 0

    # Replace in paragraphs
    for para in doc.paragraphs:
        if placeholder in para.text:
            if replace_text_in_paragraph(para, placeholder, value):
                count += 1

    # Replace in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, value)
                    count += 1

    if count == 0:
        return {"error": f"Placeholder '{placeholder}' was not found in the document."}

    save_document(doc)

    return {
        "status": "success",
        "message": f"Replaced {count} occurrence(s) of '{placeholder}' with the provided value.",
        "replacements": count
    }


@mcp.tool()
async def replace_placeholders(replacements: dict) -> dict:
    """
    TEMPLATE SUPPORT: Replace multiple placeholders at once.
    Provide a mapping of placeholder -> value pairs to fill in a template document efficiently.
    Works in both paragraphs AND table cells.

    Args:
        replacements: Object mapping placeholders to their values, e.g., {"<<Name>>": "John", "<<Date>>": "2024-01-15"}
    """
    doc = get_document()
    total_count = 0
    results = {}

    for placeholder, value in replacements.items():
        count = 0

        # Replace in paragraphs
        for para in doc.paragraphs:
            if placeholder in para.text:
                if replace_text_in_paragraph(para, placeholder, value):
                    count += 1

        # Replace in table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)
                        count += 1

        results[placeholder] = count
        total_count += count

    save_document(doc)

    return {
        "status": "success",
        "message": f"Replaced {total_count} placeholder occurrence(s) across {len(replacements)} unique placeholder(s).",
        "details": results
    }


# ============================================
# Table Tools
# ============================================

@mcp.tool()
async def insert_table(text: str, has_header: bool = True, query: str = None) -> dict:
    """
    INSERT A PROPER WORD TABLE into the document.
    Automatically detects and converts markdown or tab-delimited table text into a real Word table.

    Supports two formats:
    1. Markdown: | Column1 | Column2 |
    2. Tab-delimited: Column1\tColumn2

    Args:
        text: The table content (markdown or tab-delimited format)
        has_header: Whether the first row is a header (default True)
        query: Optional - text to search for to insert table after. If not provided, adds to end.
    """
    doc = get_document()

    # Detect and parse table format
    format_type, table_data = detect_table_format(text)

    if not format_type:
        return {
            "error": "Could not detect table format. Please use markdown (| col |) or tab-delimited format.",
            "received_text": text[:200] + "..." if len(text) > 200 else text
        }

    # If query provided, find insertion point
    insert_after_para = None
    if query:
        match = find_paragraph_by_text(doc, query, 0.5)
        if not match:
            return {"error": f"Could not find paragraph matching: '{query}'"}
        insert_after_para = match[1]

    # Create the table
    table = create_word_table(doc, table_data, has_header)

    if not table:
        return {"error": "Failed to create table"}

    # Move table to correct position if query was provided
    if insert_after_para:
        # Add space before table
        space_before = doc.add_paragraph()
        insert_after_para._element.addnext(space_before._element)

        # Add table after the space
        space_before._element.addnext(table._element)

        # Add space after table
        space_after = doc.add_paragraph()
        table._element.addnext(space_after._element)
    else:
        # If adding to end of document, add spaces too
        doc.add_paragraph()  # Space before
        # Table is already at the end
        doc.add_paragraph()  # Space after

    save_document(doc)

    return {
        "status": "success",
        "message": f"Inserted {len(table_data)}x{len(table_data[0])} table ({format_type} format detected)",
        "rows": len(table_data),
        "columns": len(table_data[0]),
        "format_detected": format_type
    }


@mcp.tool()
async def convert_text_to_table(query: str, has_header: bool = True, threshold: float = 0.5) -> dict:
    """
    CONVERT EXISTING TEXT PARAGRAPH(S) into a proper Word table.
    Finds paragraph(s) containing table-like text and replaces them with a real Word table.

    Use this when GPT has already inserted table text (markdown or tab-delimited) and you want to
    convert it to a proper formatted table.

    Args:
        query: Text to search for to find the paragraph(s) containing the table
        has_header: Whether the first row is a header (default True)
        threshold: Minimum similarity score for text matching (0-1, default 0.5)
    """
    doc = get_document()

    # Find the paragraph containing table-like text
    match = find_paragraph_by_text(doc, query, threshold)

    if not match:
        return {"error": f"No paragraph found matching: '{query}'"}

    idx, para, score = match

    # Collect consecutive paragraphs that might be part of the table
    table_paragraphs = [para]
    table_indices = [idx]

    # Check if following paragraphs are also table rows
    for i in range(idx + 1, min(idx + 20, len(doc.paragraphs))):
        next_para = doc.paragraphs[i]
        text = next_para.text.strip()

        if not text:
            break  # Empty line signals end of table

        # Check if it looks like a table row
        if '|' in text or '\t' in text:
            table_paragraphs.append(next_para)
            table_indices.append(i)
        else:
            break

    # Combine all paragraphs into one text block
    combined_text = '\n'.join(p.text for p in table_paragraphs)

    # Detect and parse table format
    format_type, table_data = detect_table_format(combined_text)

    if not format_type:
        return {
            "error": "Could not detect table format in the matched paragraph(s).",
            "matched_text": combined_text[:200] + "..." if len(combined_text) > 200 else combined_text
        }

    # Create the Word table
    table = create_word_table(doc, table_data, has_header)

    if not table:
        return {"error": "Failed to create table"}

    # Add space before table
    space_before = doc.add_paragraph()
    table_paragraphs[0]._element.addprevious(space_before._element)

    # Insert table after the space
    space_before._element.addnext(table._element)

    # Add space after table
    space_after = doc.add_paragraph()
    table._element.addnext(space_after._element)

    # Delete all the old text paragraphs
    for p in table_paragraphs:
        p._element.getparent().remove(p._element)

    save_document(doc)

    return {
        "status": "success",
        "message": f"Converted {len(table_paragraphs)} paragraph(s) to {len(table_data)}x{len(table_data[0])} table",
        "rows": len(table_data),
        "columns": len(table_data[0]),
        "format_detected": format_type,
        "paragraphs_removed": len(table_paragraphs)
    }


# ============================================
# Table Manipulation Tools
# ============================================

@mcp.tool()
async def list_tables() -> dict:
    """
    List all tables in the document with their dimensions and preview of first row.
    Returns table indices that can be used with other table tools.
    """
    doc = get_document()
    tables_info = []

    for idx, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns) if rows > 0 else 0

        # Get first row as preview
        first_row = []
        if rows > 0:
            first_row = [cell.text.strip() for cell in table.rows[0].cells]

        tables_info.append({
            "table_index": idx,
            "rows": rows,
            "columns": cols,
            "first_row_preview": first_row[:5]  # First 5 cells
        })

    return {
        "total_tables": len(tables_info),
        "tables": tables_info
    }


@mcp.tool()
async def read_table(table_index: int) -> dict:
    """
    Read the complete contents of a table by its index.

    Args:
        table_index: The index of the table (from list_tables, starting at 0)
    """
    doc = get_document()

    if table_index < 0 or table_index >= len(doc.tables):
        return {"error": f"Table index {table_index} not found. Document has {len(doc.tables)} table(s)."}

    table = doc.tables[table_index]
    table_data = []

    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        table_data.append(row_data)

    return {
        "table_index": table_index,
        "rows": len(table.rows),
        "columns": len(table.columns),
        "data": table_data
    }


@mcp.tool()
async def update_table_cell(table_index: int, row: int, column: int, text: str) -> dict:
    """
    Update a specific cell in a table by row and column index.
    This is the PRIMARY tool for editing table cells directly.

    Args:
        table_index: The index of the table (from list_tables, starting at 0)
        row: The row index (starting at 0)
        column: The column index (starting at 0)
        text: The new text content for the cell
    """
    doc = get_document()

    if table_index < 0 or table_index >= len(doc.tables):
        return {"error": f"Table index {table_index} not found. Document has {len(doc.tables)} table(s)."}

    table = doc.tables[table_index]

    if row < 0 or row >= len(table.rows):
        return {"error": f"Row index {row} out of range. Table has {len(table.rows)} rows."}

    if column < 0 or column >= len(table.columns):
        return {"error": f"Column index {column} out of range. Table has {len(table.columns)} columns."}

    # Update the cell
    cell = table.rows[row].cells[column]
    cell.text = text

    save_document(doc)

    return {
        "status": "success",
        "message": f"Updated cell at table {table_index}, row {row}, column {column}",
        "new_value": text
    }


@mcp.tool()
async def add_table_row(table_index: int, row_data: list = None) -> dict:
    """
    Add a new row to the end of a table, optionally with data.

    Args:
        table_index: The index of the table (from list_tables, starting at 0)
        row_data: Optional list of cell values for the new row
    """
    doc = get_document()

    if table_index < 0 or table_index >= len(doc.tables):
        return {"error": f"Table index {table_index} not found. Document has {len(doc.tables)} table(s)."}

    table = doc.tables[table_index]
    new_row = table.add_row()

    # Fill in data if provided
    if row_data:
        for col_idx, value in enumerate(row_data):
            if col_idx < len(new_row.cells):
                new_row.cells[col_idx].text = str(value)

    save_document(doc)

    return {
        "status": "success",
        "message": f"Added row to table {table_index}",
        "new_row_index": len(table.rows) - 1,
        "total_rows": len(table.rows)
    }


@mcp.tool()
async def update_table_row(table_index: int, row: int, row_data: list) -> dict:
    """
    Update an entire row in a table with new data.

    Args:
        table_index: The index of the table (from list_tables, starting at 0)
        row: The row index (starting at 0)
        row_data: List of cell values for the row
    """
    doc = get_document()

    if table_index < 0 or table_index >= len(doc.tables):
        return {"error": f"Table index {table_index} not found. Document has {len(doc.tables)} table(s)."}

    table = doc.tables[table_index]

    if row < 0 or row >= len(table.rows):
        return {"error": f"Row index {row} out of range. Table has {len(table.rows)} rows."}

    # Update each cell in the row
    table_row = table.rows[row]
    for col_idx, value in enumerate(row_data):
        if col_idx < len(table_row.cells):
            table_row.cells[col_idx].text = str(value)

    save_document(doc)

    return {
        "status": "success",
        "message": f"Updated row {row} in table {table_index}",
        "cells_updated": min(len(row_data), len(table_row.cells))
    }


@mcp.tool()
async def delete_table_row(table_index: int, row: int) -> dict:
    """
    Delete a row from a table.

    Args:
        table_index: The index of the table (from list_tables, starting at 0)
        row: The row index to delete (starting at 0)
    """
    doc = get_document()

    if table_index < 0 or table_index >= len(doc.tables):
        return {"error": f"Table index {table_index} not found. Document has {len(doc.tables)} table(s)."}

    table = doc.tables[table_index]

    if row < 0 or row >= len(table.rows):
        return {"error": f"Row index {row} out of range. Table has {len(table.rows)} rows."}

    # Delete the row using XML manipulation
    table._element.remove(table.rows[row]._element)

    save_document(doc)

    return {
        "status": "success",
        "message": f"Deleted row {row} from table {table_index}",
        "remaining_rows": len(table.rows)
    }


# ============================================
# Formatting Tools
# ============================================

@mcp.tool()
async def format_paragraph(
    query: str,
    bold: bool = None,
    italic: bool = None,
    underline: bool = None,
    alignment: str = None,
    font_size: int = None,
    style: str = None,
    threshold: float = 0.5
) -> dict:
    """
    Apply formatting to an existing paragraph found by fuzzy text search.
    Use this to format paragraphs that have already been added to the document.

    Args:
        query: Text to search for to find the paragraph (fuzzy matching)
        bold: Make text bold (True/False, None = don't change)
        italic: Make text italic (True/False, None = don't change)
        underline: Make text underlined (True/False, None = don't change)
        alignment: Text alignment - "left", "center", "right", "justify" (None = don't change)
        font_size: Font size in points, e.g., 12, 14, 16 (None = don't change)
        style: Word style name - "Heading 1", "Title", etc. (None = don't change)
        threshold: Minimum similarity score for text matching (0-1, default 0.5)
    """
    doc = get_document()

    # Find the paragraph
    match = find_paragraph_by_text(doc, query, threshold)

    if not match:
        return {"error": f"No paragraph found matching: '{query}'"}

    idx, para, score = match

    # Apply text formatting (bold, italic, underline)
    if bold is not None or italic is not None or underline is not None:
        for run in para.runs:
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic
            if underline is not None:
                run.underline = underline

    # Apply font size
    if font_size:
        for run in para.runs:
            run.font.size = Pt(font_size)

    # Apply alignment
    if alignment:
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        if alignment.lower() in alignment_map:
            para.alignment = alignment_map[alignment.lower()]

    # Apply style
    if style:
        try:
            para.style = style
        except KeyError:
            pass  # Style doesn't exist, ignore

    save_document(doc)

    return {
        "status": "success",
        "message": f"Applied formatting to paragraph {idx} (match score: {score:.2f})",
        "matched_text": para.text[:100] + "..." if len(para.text) > 100 else para.text
    }


# ============================================
# Text Replacement Tool
# ============================================

@mcp.tool()
async def replace_text(old_text: str, new_text: str) -> dict:
    """
    FIND AND REPLACE any text anywhere in the document.
    This is useful for editing generated content without replacing entire paragraphs.

    Use this tool to:
    - Fix typos or errors in previously generated text
    - Update specific words or phrases throughout the document
    - Make surgical edits to GPT-generated content

    Args:
        old_text: The exact text to find and replace
        new_text: The text to replace it with
    """
    doc = get_document()
    count = 0
    affected_paragraphs = []

    for idx, para in enumerate(doc.paragraphs):
        if old_text in para.text:
            if replace_text_in_paragraph(para, old_text, new_text):
                count += 1
                affected_paragraphs.append({
                    "id": f"para-{idx}",
                    "preview": para.text[:100] + "..." if len(para.text) > 100 else para.text
                })

    if count == 0:
        return {
            "error": f"Text '{old_text}' was not found in the document.",
            "suggestion": "Make sure the text matches exactly (case-sensitive)."
        }

    save_document(doc)

    return {
        "status": "success",
        "message": f"Replaced {count} occurrence(s) of '{old_text}' with '{new_text}'.",
        "replacements": count,
        "affected_paragraphs": affected_paragraphs
    }


# ============================================
# Save Tool
# ============================================

@mcp.tool()
async def save_document_as(filename: str) -> dict:
    """
    Save the document to a new filename.
    Use this to create a copy or save with a different name.

    Args:
        filename: The filename to save as (e.g., "output.docx")
    """
    doc = get_document()

    # Ensure .docx extension
    if not filename.endswith('.docx'):
        filename += '.docx'

    doc.save(filename)

    return {
        "status": "success",
        "message": f"Document saved to: {filename}",
        "path": os.path.abspath(filename)
    }


# ============================================
# Paragraph Management Tools
# ============================================

@mcp.tool()
async def delete_paragraph(query: str, threshold: float = 0.5) -> dict:
    """
    DELETE a paragraph from the document by fuzzy text search.

    Args:
        query: The text to search for - finds and deletes the paragraph containing or matching this text
        threshold: Minimum similarity score 0-1 (default 0.5)
    """
    doc = get_document()

    match = find_paragraph_by_text(doc, query, threshold)

    if not match:
        return {"error": f"No paragraph found matching: '{query}'"}

    idx, para, score = match
    deleted_text = para.text

    # Delete the paragraph using XML manipulation
    para._element.getparent().remove(para._element)

    save_document(doc)

    return {
        "status": "success",
        "message": f"Deleted paragraph {idx} (match score: {score:.2f})",
        "deleted_text": deleted_text[:100] + "..." if len(deleted_text) > 100 else deleted_text
    }


@mcp.tool()
async def move_paragraph(query: str, target_query: str, position: str = "after", threshold: float = 0.5) -> dict:
    """
    MOVE a paragraph to a new location in the document.

    Args:
        query: Text to search for to find the paragraph to move
        target_query: Text to search for to find the target location
        position: Where to place relative to target - "before" or "after" (default "after")
        threshold: Minimum similarity score 0-1 (default 0.5)
    """
    doc = get_document()

    # Find the source paragraph
    source_match = find_paragraph_by_text(doc, query, threshold)
    if not source_match:
        return {"error": f"No paragraph found matching source: '{query}'"}

    source_idx, source_para, source_score = source_match

    # Find the target paragraph
    target_match = find_paragraph_by_text(doc, target_query, threshold)
    if not target_match:
        return {"error": f"No paragraph found matching target: '{target_query}'"}

    target_idx, target_para, target_score = target_match

    if source_idx == target_idx:
        return {"error": "Source and target paragraphs are the same"}

    # Store the source element
    source_element = source_para._element

    # Remove from current position
    source_element.getparent().remove(source_element)

    # Insert at new position
    if position.lower() == "before":
        target_para._element.addprevious(source_element)
    else:
        target_para._element.addnext(source_element)

    save_document(doc)

    return {
        "status": "success",
        "message": f"Moved paragraph {position} target",
        "moved_text": source_para.text[:100] + "..." if len(source_para.text) > 100 else source_para.text,
        "target_text": target_para.text[:100] + "..." if len(target_para.text) > 100 else target_para.text
    }


@mcp.tool()
async def merge_paragraphs(query1: str, query2: str, separator: str = " ", threshold: float = 0.5) -> dict:
    """
    MERGE two paragraphs into one. The second paragraph's text is appended to the first,
    then the second paragraph is deleted.

    Args:
        query1: Text to search for the first paragraph (will be kept)
        query2: Text to search for the second paragraph (will be merged and deleted)
        separator: Text to insert between the two paragraphs (default is a space)
        threshold: Minimum similarity score 0-1 (default 0.5)
    """
    doc = get_document()

    # Find both paragraphs
    match1 = find_paragraph_by_text(doc, query1, threshold)
    if not match1:
        return {"error": f"No paragraph found matching: '{query1}'"}

    match2 = find_paragraph_by_text(doc, query2, threshold)
    if not match2:
        return {"error": f"No paragraph found matching: '{query2}'"}

    idx1, para1, score1 = match1
    idx2, para2, score2 = match2

    if idx1 == idx2:
        return {"error": "Both queries match the same paragraph"}

    # Merge text
    original_text1 = para1.text
    para2_text = para2.text

    # Add separator and second paragraph text to first paragraph
    if para1.runs:
        para1.runs[-1].text += separator + para2_text
    else:
        para1.add_run(separator + para2_text)

    # Delete the second paragraph
    para2._element.getparent().remove(para2._element)

    save_document(doc)

    return {
        "status": "success",
        "message": f"Merged paragraphs {idx1} and {idx2}",
        "merged_text": para1.text[:150] + "..." if len(para1.text) > 150 else para1.text
    }


# ============================================
# Document Structure & Stats Tools
# ============================================

@mcp.tool()
async def get_document_outline() -> dict:
    """
    Get the document's heading structure/outline.
    Returns a hierarchical view of all headings in the document.
    Useful for navigating large documents.
    """
    doc = get_document()
    outline = []

    for idx, para in enumerate(doc.paragraphs):
        if para.style and para.style.name.startswith('Heading'):
            # Extract heading level
            try:
                level = int(para.style.name.replace('Heading ', ''))
            except ValueError:
                level = 0

            outline.append({
                "id": f"para-{idx}",
                "level": level,
                "text": para.text[:100] + "..." if len(para.text) > 100 else para.text,
                "style": para.style.name
            })
        elif para.style and para.style.name == 'Title':
            outline.append({
                "id": f"para-{idx}",
                "level": 0,
                "text": para.text[:100] + "..." if len(para.text) > 100 else para.text,
                "style": "Title"
            })

    return {
        "total_headings": len(outline),
        "outline": outline
    }


@mcp.tool()
async def get_document_stats() -> dict:
    """
    Get document statistics including word count, character count, paragraph count, etc.
    """
    doc = get_document()

    total_paragraphs = 0
    non_empty_paragraphs = 0
    total_words = 0
    total_characters = 0
    total_characters_no_spaces = 0
    total_tables = len(doc.tables)
    total_images = 0

    for para in doc.paragraphs:
        total_paragraphs += 1
        text = para.text.strip()
        if text:
            non_empty_paragraphs += 1
            words = text.split()
            total_words += len(words)
            total_characters += len(text)
            total_characters_no_spaces += len(text.replace(' ', ''))

    # Count images (inline shapes)
    for para in doc.paragraphs:
        for run in para.runs:
            if run._element.xpath('.//a:blip'):
                total_images += 1

    # Count table cells content
    table_words = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_words += len(cell.text.split())

    return {
        "paragraphs": {
            "total": total_paragraphs,
            "non_empty": non_empty_paragraphs
        },
        "words": total_words + table_words,
        "characters": total_characters,
        "characters_no_spaces": total_characters_no_spaces,
        "tables": total_tables,
        "images": total_images
    }


# ============================================
# Page & Section Tools
# ============================================

@mcp.tool()
async def insert_page_break(query: str = None, threshold: float = 0.5) -> dict:
    """
    Insert a page break in the document.

    Args:
        query: Optional text to search for - inserts page break after matching paragraph.
               If not provided, adds page break at the end of the document.
        threshold: Minimum similarity score 0-1 (default 0.5)
    """
    doc = get_document()

    if query:
        match = find_paragraph_by_text(doc, query, threshold)
        if not match:
            return {"error": f"No paragraph found matching: '{query}'"}

        idx, para, score = match

        # Add a new paragraph with page break after the matched paragraph
        new_para = doc.add_paragraph()
        run = new_para.add_run()
        run.add_break(WD_BREAK.PAGE)

        # Move it after the matched paragraph
        para._element.addnext(new_para._element)

        save_document(doc)

        return {
            "status": "success",
            "message": f"Page break inserted after paragraph {idx}",
            "after_text": para.text[:100] + "..." if len(para.text) > 100 else para.text
        }
    else:
        # Add page break at the end
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_break(WD_BREAK.PAGE)

        save_document(doc)

        return {
            "status": "success",
            "message": "Page break added at the end of the document"
        }


@mcp.tool()
async def insert_image(
    image_path: str,
    query: str = None,
    width: float = None,
    height: float = None,
    threshold: float = 0.5
) -> dict:
    """
    Insert an image into the document.

    Args:
        image_path: Path to the image file (supports PNG, JPG, GIF, BMP, etc.)
        query: Optional text to search for - inserts image after matching paragraph.
               If not provided, adds image at the end of the document.
        width: Optional width in inches (e.g., 4.0 for 4 inches)
        height: Optional height in inches (e.g., 3.0 for 3 inches)
        threshold: Minimum similarity score 0-1 (default 0.5)
    """
    # Check if image exists
    if not os.path.exists(image_path):
        return {"error": f"Image file not found: {image_path}"}

    doc = get_document()

    # Prepare size arguments
    size_kwargs = {}
    if width:
        size_kwargs['width'] = Inches(width)
    if height:
        size_kwargs['height'] = Inches(height)

    if query:
        match = find_paragraph_by_text(doc, query, threshold)
        if not match:
            return {"error": f"No paragraph found matching: '{query}'"}

        idx, para, score = match

        # Create a new paragraph for the image
        new_para = doc.add_paragraph()
        run = new_para.add_run()
        run.add_picture(image_path, **size_kwargs)

        # Move it after the matched paragraph
        para._element.addnext(new_para._element)

        save_document(doc)

        return {
            "status": "success",
            "message": f"Image inserted after paragraph {idx}",
            "image": os.path.basename(image_path),
            "after_text": para.text[:100] + "..." if len(para.text) > 100 else para.text
        }
    else:
        # Add image at the end
        doc.add_picture(image_path, **size_kwargs)

        save_document(doc)

        return {
            "status": "success",
            "message": "Image added at the end of the document",
            "image": os.path.basename(image_path)
        }


# ============================================
# List Tools
# ============================================

@mcp.tool()
async def create_list(
    items: list,
    list_type: str = "bullet",
    query: str = None,
    threshold: float = 0.5
) -> dict:
    """
    Create a bulleted or numbered list in the document.

    Args:
        items: List of strings - each string becomes a list item
        list_type: Type of list - "bullet" or "number" (default "bullet")
        query: Optional text to search for - inserts list after matching paragraph.
               If not provided, adds list at the end of the document.
        threshold: Minimum similarity score 0-1 (default 0.5)
    """
    if not items:
        return {"error": "No items provided for the list"}

    doc = get_document()

    # Determine the list style
    if list_type.lower() in ["number", "numbered", "ordered"]:
        style_name = "List Number"
    else:
        style_name = "List Bullet"

    insert_after_para = None
    if query:
        match = find_paragraph_by_text(doc, query, threshold)
        if not match:
            return {"error": f"No paragraph found matching: '{query}'"}
        idx, insert_after_para, score = match

    # Create list items
    created_paragraphs = []
    for item in items:
        para = doc.add_paragraph(item)
        try:
            para.style = style_name
        except KeyError:
            # Style doesn't exist, just leave as normal paragraph
            pass
        created_paragraphs.append(para)

    # Move paragraphs to correct position if query was provided
    if insert_after_para:
        # Move in reverse order to maintain order
        for para in reversed(created_paragraphs):
            insert_after_para._element.addnext(para._element)

    save_document(doc)

    return {
        "status": "success",
        "message": f"Created {list_type} list with {len(items)} items",
        "items_count": len(items),
        "list_type": list_type
    }


@mcp.tool()
async def add_list_item(
    query: str,
    text: str,
    list_type: str = "bullet",
    threshold: float = 0.5
) -> dict:
    """
    Add an item to an existing list. Finds a list item and adds a new item after it.

    Args:
        query: Text to search for to find existing list item
        text: The text for the new list item
        list_type: Type of list - "bullet" or "number" (default "bullet")
        threshold: Minimum similarity score 0-1 (default 0.5)
    """
    doc = get_document()

    match = find_paragraph_by_text(doc, query, threshold)
    if not match:
        return {"error": f"No paragraph found matching: '{query}'"}

    idx, para, score = match

    # Determine the list style
    if list_type.lower() in ["number", "numbered", "ordered"]:
        style_name = "List Number"
    else:
        style_name = "List Bullet"

    # Create new list item
    new_para = doc.add_paragraph(text)
    try:
        new_para.style = style_name
    except KeyError:
        pass

    # Insert after the matched paragraph
    para._element.addnext(new_para._element)

    save_document(doc)

    return {
        "status": "success",
        "message": f"Added list item after paragraph {idx}",
        "new_item": text
    }


# ============================================
# Advanced Formatting Tools
# ============================================

@mcp.tool()
async def clear_formatting(query: str, threshold: float = 0.5) -> dict:
    """
    Remove ALL formatting from a paragraph, resetting it to plain text with Normal style.

    Args:
        query: Text to search for to find the paragraph
        threshold: Minimum similarity score 0-1 (default 0.5)
    """
    doc = get_document()

    match = find_paragraph_by_text(doc, query, threshold)
    if not match:
        return {"error": f"No paragraph found matching: '{query}'"}

    idx, para, score = match

    # Reset paragraph style to Normal
    try:
        para.style = 'Normal'
    except KeyError:
        pass

    # Clear run-level formatting
    for run in para.runs:
        run.bold = False
        run.italic = False
        run.underline = False
        run.font.size = None  # Reset to default
        run.font.color.rgb = None  # Reset color
        run.font.name = None  # Reset font

    # Reset paragraph alignment
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    save_document(doc)

    return {
        "status": "success",
        "message": f"Cleared all formatting from paragraph {idx}",
        "text": para.text[:100] + "..." if len(para.text) > 100 else para.text
    }


# ============================================
# Hyperlink & Bookmark Tools
# ============================================

def add_hyperlink(paragraph, url, text):
    """
    Helper function to add a hyperlink to a paragraph.
    python-docx doesn't have built-in hyperlink support, so we use XML manipulation.
    """
    # Get the document part
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run for the hyperlink text
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add hyperlink styling (blue and underlined)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    # Add the text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


@mcp.tool()
async def insert_hyperlink(
    url: str,
    display_text: str,
    query: str = None,
    threshold: float = 0.5
) -> dict:
    """
    Insert a clickable hyperlink into the document.

    Args:
        url: The URL the hyperlink points to (e.g., "https://example.com")
        display_text: The visible text for the hyperlink
        query: Optional text to search for - adds hyperlink to end of matching paragraph.
               If not provided, creates a new paragraph with the hyperlink.
        threshold: Minimum similarity score 0-1 (default 0.5)
    """
    doc = get_document()

    if query:
        match = find_paragraph_by_text(doc, query, threshold)
        if not match:
            return {"error": f"No paragraph found matching: '{query}'"}

        idx, para, score = match

        # Add hyperlink to the paragraph
        add_hyperlink(para, url, display_text)

        save_document(doc)

        return {
            "status": "success",
            "message": f"Hyperlink added to paragraph {idx}",
            "url": url,
            "display_text": display_text
        }
    else:
        # Create new paragraph with hyperlink
        para = doc.add_paragraph()
        add_hyperlink(para, url, display_text)

        save_document(doc)

        return {
            "status": "success",
            "message": "Hyperlink added in new paragraph at end of document",
            "url": url,
            "display_text": display_text
        }


@mcp.tool()
async def insert_bookmark(query: str, bookmark_name: str, threshold: float = 0.5) -> dict:
    """
    Insert a bookmark at a specific paragraph location.
    Bookmarks can be used for internal document references.

    Args:
        query: Text to search for to find the paragraph to bookmark
        bookmark_name: Name for the bookmark (no spaces, use underscores)
        threshold: Minimum similarity score 0-1 (default 0.5)
    """
    doc = get_document()

    match = find_paragraph_by_text(doc, query, threshold)
    if not match:
        return {"error": f"No paragraph found matching: '{query}'"}

    idx, para, score = match

    # Clean bookmark name (remove spaces, special characters)
    clean_name = re.sub(r'[^a-zA-Z0-9_]', '_', bookmark_name)

    # Create bookmark start element
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(idx))  # Use paragraph index as bookmark ID
    bookmark_start.set(qn('w:name'), clean_name)

    # Create bookmark end element
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), str(idx))

    # Insert bookmark around paragraph content
    para._p.insert(0, bookmark_start)
    para._p.append(bookmark_end)

    save_document(doc)

    return {
        "status": "success",
        "message": f"Bookmark '{clean_name}' added to paragraph {idx}",
        "bookmark_name": clean_name,
        "paragraph_text": para.text[:100] + "..." if len(para.text) > 100 else para.text
    }


# ============================================
# Header & Footer Tools
# ============================================

@mcp.tool()
async def get_header(section_index: int = 0) -> dict:
    """
    Get the header content from a document section.

    Args:
        section_index: Index of the section (default 0 for first section)
    """
    doc = get_document()

    if section_index >= len(doc.sections):
        return {"error": f"Section {section_index} not found. Document has {len(doc.sections)} section(s)."}

    section = doc.sections[section_index]
    header = section.header

    # Get all paragraphs in header
    header_text = []
    for para in header.paragraphs:
        if para.text.strip():
            header_text.append(para.text)

    return {
        "section_index": section_index,
        "has_header": len(header_text) > 0,
        "header_text": "\n".join(header_text) if header_text else "(empty)",
        "paragraphs": len(header.paragraphs)
    }


@mcp.tool()
async def set_header(
    text: str,
    section_index: int = 0,
    alignment: str = "center"
) -> dict:
    """
    Set the header content for a document section.

    Args:
        text: The header text to set
        section_index: Index of the section (default 0 for first section)
        alignment: Text alignment - "left", "center", "right" (default "center")
    """
    doc = get_document()

    if section_index >= len(doc.sections):
        return {"error": f"Section {section_index} not found. Document has {len(doc.sections)} section(s)."}

    section = doc.sections[section_index]
    header = section.header

    # Clear existing header content
    for para in header.paragraphs:
        para.clear()

    # Add new header text
    if header.paragraphs:
        para = header.paragraphs[0]
    else:
        para = header.add_paragraph()

    para.text = text

    # Apply alignment
    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT
    }
    if alignment.lower() in alignment_map:
        para.alignment = alignment_map[alignment.lower()]

    save_document(doc)

    return {
        "status": "success",
        "message": f"Header set for section {section_index}",
        "header_text": text
    }


@mcp.tool()
async def get_footer(section_index: int = 0) -> dict:
    """
    Get the footer content from a document section.

    Args:
        section_index: Index of the section (default 0 for first section)
    """
    doc = get_document()

    if section_index >= len(doc.sections):
        return {"error": f"Section {section_index} not found. Document has {len(doc.sections)} section(s)."}

    section = doc.sections[section_index]
    footer = section.footer

    # Get all paragraphs in footer
    footer_text = []
    for para in footer.paragraphs:
        if para.text.strip():
            footer_text.append(para.text)

    return {
        "section_index": section_index,
        "has_footer": len(footer_text) > 0,
        "footer_text": "\n".join(footer_text) if footer_text else "(empty)",
        "paragraphs": len(footer.paragraphs)
    }


@mcp.tool()
async def set_footer(
    text: str,
    section_index: int = 0,
    alignment: str = "center"
) -> dict:
    """
    Set the footer content for a document section.

    Args:
        text: The footer text to set
        section_index: Index of the section (default 0 for first section)
        alignment: Text alignment - "left", "center", "right" (default "center")
    """
    doc = get_document()

    if section_index >= len(doc.sections):
        return {"error": f"Section {section_index} not found. Document has {len(doc.sections)} section(s)."}

    section = doc.sections[section_index]
    footer = section.footer

    # Clear existing footer content
    for para in footer.paragraphs:
        para.clear()

    # Add new footer text
    if footer.paragraphs:
        para = footer.paragraphs[0]
    else:
        para = footer.add_paragraph()

    para.text = text

    # Apply alignment
    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT
    }
    if alignment.lower() in alignment_map:
        para.alignment = alignment_map[alignment.lower()]

    save_document(doc)

    return {
        "status": "success",
        "message": f"Footer set for section {section_index}",
        "footer_text": text
    }


# ============================================
# Document Properties Tools
# ============================================

@mcp.tool()
async def get_document_properties() -> dict:
    """
    Get document metadata/properties like title, author, subject, etc.
    """
    doc = get_document()
    props = doc.core_properties

    return {
        "title": props.title or "(not set)",
        "author": props.author or "(not set)",
        "subject": props.subject or "(not set)",
        "keywords": props.keywords or "(not set)",
        "comments": props.comments or "(not set)",
        "category": props.category or "(not set)",
        "created": str(props.created) if props.created else "(not set)",
        "modified": str(props.modified) if props.modified else "(not set)",
        "last_modified_by": props.last_modified_by or "(not set)"
    }


@mcp.tool()
async def set_document_properties(
    title: str = None,
    author: str = None,
    subject: str = None,
    keywords: str = None,
    comments: str = None,
    category: str = None
) -> dict:
    """
    Set document metadata/properties.
    Only provided fields will be updated; others remain unchanged.

    Args:
        title: Document title
        author: Document author name
        subject: Document subject
        keywords: Keywords (comma-separated)
        comments: Comments/description
        category: Document category
    """
    doc = get_document()
    props = doc.core_properties

    updated = []

    if title is not None:
        props.title = title
        updated.append("title")
    if author is not None:
        props.author = author
        updated.append("author")
    if subject is not None:
        props.subject = subject
        updated.append("subject")
    if keywords is not None:
        props.keywords = keywords
        updated.append("keywords")
    if comments is not None:
        props.comments = comments
        updated.append("comments")
    if category is not None:
        props.category = category
        updated.append("category")

    if not updated:
        return {"error": "No properties provided to update"}

    save_document(doc)

    return {
        "status": "success",
        "message": f"Updated document properties: {', '.join(updated)}",
        "updated_fields": updated
    }


# ============================================
# Table Column Tools
# ============================================

@mcp.tool()
async def add_table_column(
    table_index: int,
    column_data: list = None,
    position: int = None
) -> dict:
    """
    Add a new column to a table.

    Args:
        table_index: The index of the table (from list_tables, starting at 0)
        column_data: Optional list of cell values for the new column (one per row)
        position: Optional column position (0-indexed). If not provided, adds at the end.
    """
    doc = get_document()

    if table_index < 0 or table_index >= len(doc.tables):
        return {"error": f"Table index {table_index} not found. Document has {len(doc.tables)} table(s)."}

    table = doc.tables[table_index]
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    # Determine insertion position
    if position is None:
        position = num_cols
    elif position < 0 or position > num_cols:
        position = num_cols

    # Add a cell to each row
    for row_idx, row in enumerate(table.rows):
        # Create a new cell element
        new_cell = OxmlElement('w:tc')

        # Add cell properties
        tcPr = OxmlElement('w:tcPr')
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), '0')
        tcW.set(qn('w:type'), 'auto')
        tcPr.append(tcW)
        new_cell.append(tcPr)

        # Add paragraph to cell
        p = OxmlElement('w:p')
        new_cell.append(p)

        # Insert at the correct position
        cells = row._tr.findall(qn('w:tc'))
        if position < len(cells):
            cells[position].addprevious(new_cell)
        else:
            row._tr.append(new_cell)

    # Reload the table to get proper cell references
    table = doc.tables[table_index]

    # Fill in data if provided
    if column_data:
        for row_idx, value in enumerate(column_data):
            if row_idx < len(table.rows):
                row = table.rows[row_idx]
                if position < len(row.cells):
                    row.cells[position].text = str(value)

    save_document(doc)

    return {
        "status": "success",
        "message": f"Added column at position {position} in table {table_index}",
        "new_column_count": num_cols + 1
    }


@mcp.tool()
async def delete_table_column(table_index: int, column: int) -> dict:
    """
    Delete a column from a table.

    Args:
        table_index: The index of the table (from list_tables, starting at 0)
        column: The column index to delete (starting at 0)
    """
    doc = get_document()

    if table_index < 0 or table_index >= len(doc.tables):
        return {"error": f"Table index {table_index} not found. Document has {len(doc.tables)} table(s)."}

    table = doc.tables[table_index]

    if column < 0 or column >= len(table.columns):
        return {"error": f"Column index {column} out of range. Table has {len(table.columns)} columns."}

    # Delete the cell at the specified column in each row
    for row in table.rows:
        cells = row._tr.findall(qn('w:tc'))
        if column < len(cells):
            row._tr.remove(cells[column])

    save_document(doc)

    return {
        "status": "success",
        "message": f"Deleted column {column} from table {table_index}",
        "remaining_columns": len(table.columns) - 1
    }


# ============================================
# Document Creation Tools
# ============================================

@mcp.tool()
async def create_document(filename: str, title: str = None, switch_to: bool = True) -> dict:
    """
    Create a new blank DOCX document.

    Args:
        filename: Name for the new document (e.g., "report.docx")
        title: Optional title to add as the first paragraph
        switch_to: Whether to switch to the new document after creation (default True)
    """
    global CURRENT_DOCX_PATH

    # Ensure .docx extension
    if not filename.endswith('.docx'):
        filename += '.docx'

    # Check if file already exists
    if os.path.exists(filename):
        return {"error": f"File '{filename}' already exists. Use a different name or delete the existing file."}

    # Create new document
    doc = Document()

    # Add title if provided
    if title:
        para = doc.add_paragraph(title)
        para.style = 'Title'

    # Save the document
    doc.save(filename)

    # Switch to the new document if requested
    if switch_to:
        CURRENT_DOCX_PATH = filename

    return {
        "status": "success",
        "message": f"Created new document: {filename}",
        "path": os.path.abspath(filename),
        "switched_to": switch_to
    }


# ============================================
# Table Management Tools
# ============================================

@mcp.tool()
async def delete_table(table_index: int) -> dict:
    """
    Delete an entire table from the document.

    Args:
        table_index: The index of the table to delete (from list_tables, starting at 0)
    """
    doc = get_document()

    if table_index < 0 or table_index >= len(doc.tables):
        return {"error": f"Table index {table_index} not found. Document has {len(doc.tables)} table(s)."}

    table = doc.tables[table_index]

    # Get table info before deleting
    rows = len(table.rows)
    cols = len(table.columns)

    # Delete the table using XML manipulation
    table._element.getparent().remove(table._element)

    save_document(doc)

    return {
        "status": "success",
        "message": f"Deleted table {table_index} ({rows}x{cols})",
        "remaining_tables": len(doc.tables) - 1
    }


@mcp.tool()
async def merge_table_cells(
    table_index: int,
    start_row: int,
    start_col: int,
    end_row: int,
    end_col: int
) -> dict:
    """
    Merge a range of cells in a table.

    Args:
        table_index: The index of the table (from list_tables, starting at 0)
        start_row: Starting row index (0-indexed)
        start_col: Starting column index (0-indexed)
        end_row: Ending row index (0-indexed, inclusive)
        end_col: Ending column index (0-indexed, inclusive)
    """
    doc = get_document()

    if table_index < 0 or table_index >= len(doc.tables):
        return {"error": f"Table index {table_index} not found. Document has {len(doc.tables)} table(s)."}

    table = doc.tables[table_index]

    # Validate indices
    if start_row < 0 or end_row >= len(table.rows):
        return {"error": f"Row indices out of range. Table has {len(table.rows)} rows."}
    if start_col < 0 or end_col >= len(table.columns):
        return {"error": f"Column indices out of range. Table has {len(table.columns)} columns."}
    if start_row > end_row or start_col > end_col:
        return {"error": "Start indices must be less than or equal to end indices."}

    # Get the cells to merge
    start_cell = table.rows[start_row].cells[start_col]
    end_cell = table.rows[end_row].cells[end_col]

    # Merge the cells
    start_cell.merge(end_cell)

    save_document(doc)

    return {
        "status": "success",
        "message": f"Merged cells from ({start_row},{start_col}) to ({end_row},{end_col}) in table {table_index}"
    }


# ============================================
# Paragraph Duplication Tools
# ============================================

@mcp.tool()
async def duplicate_paragraph(
    query: str,
    target_query: str = None,
    position: str = "after",
    threshold: float = 0.5
) -> dict:
    """
    Duplicate/copy a paragraph to another location in the document.

    Args:
        query: Text to search for to find the paragraph to duplicate
        target_query: Text to search for to find where to place the copy.
                      If not provided, places copy right after the original.
        position: Where to place relative to target - "before" or "after" (default "after")
        threshold: Minimum similarity score 0-1 (default 0.5)
    """
    doc = get_document()

    # Find the source paragraph
    source_match = find_paragraph_by_text(doc, query, threshold)
    if not source_match:
        return {"error": f"No paragraph found matching: '{query}'"}

    source_idx, source_para, source_score = source_match

    # Determine target location
    if target_query:
        target_match = find_paragraph_by_text(doc, target_query, threshold)
        if not target_match:
            return {"error": f"No paragraph found matching target: '{target_query}'"}
        target_idx, target_para, target_score = target_match
    else:
        target_para = source_para
        position = "after"

    # Create a copy of the paragraph
    new_para = doc.add_paragraph()

    # Copy the text and formatting
    for run in source_para.runs:
        new_run = new_para.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        if run.font.size:
            new_run.font.size = run.font.size
        if run.font.name:
            new_run.font.name = run.font.name

    # Copy paragraph style
    if source_para.style:
        try:
            new_para.style = source_para.style
        except:
            pass

    # Copy alignment
    if source_para.alignment:
        new_para.alignment = source_para.alignment

    # Move to correct position
    if position.lower() == "before":
        target_para._element.addprevious(new_para._element)
    else:
        target_para._element.addnext(new_para._element)

    save_document(doc)

    return {
        "status": "success",
        "message": f"Duplicated paragraph {position} target",
        "duplicated_text": source_para.text[:100] + "..." if len(source_para.text) > 100 else source_para.text
    }


@mcp.tool()
async def split_paragraph(query: str, split_at: str, threshold: float = 0.5) -> dict:
    """
    Split a paragraph into two at a specific text point.

    Args:
        query: Text to search for to find the paragraph to split
        split_at: The text at which to split the paragraph (everything after this becomes the new paragraph)
        threshold: Minimum similarity score 0-1 (default 0.5)
    """
    doc = get_document()

    match = find_paragraph_by_text(doc, query, threshold)
    if not match:
        return {"error": f"No paragraph found matching: '{query}'"}

    idx, para, score = match

    # Find the split point
    full_text = para.text
    split_index = full_text.find(split_at)

    if split_index == -1:
        return {"error": f"Split text '{split_at}' not found in the paragraph."}

    # Calculate where to split (after the split_at text)
    split_point = split_index + len(split_at)
    first_part = full_text[:split_point]
    second_part = full_text[split_point:].lstrip()

    if not second_part:
        return {"error": "Nothing to split - the split point is at the end of the paragraph."}

    # Update the original paragraph with first part
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = first_part
    else:
        para.add_run(first_part)

    # Create new paragraph with second part
    new_para = doc.add_paragraph(second_part)

    # Copy style from original
    if para.style:
        try:
            new_para.style = para.style
        except:
            pass

    # Move new paragraph right after original
    para._element.addnext(new_para._element)

    save_document(doc)

    return {
        "status": "success",
        "message": f"Split paragraph {idx} at '{split_at}'",
        "first_part": first_part[:100] + "..." if len(first_part) > 100 else first_part,
        "second_part": second_part[:100] + "..." if len(second_part) > 100 else second_part
    }


# ============================================
# Style Tools
# ============================================

@mcp.tool()
async def get_styles(style_type: str = "all") -> dict:
    """
    List all available styles in the document.

    Args:
        style_type: Type of styles to list - "paragraph", "character", "table", or "all" (default "all")
    """
    doc = get_document()
    styles_list = []

    type_map = {
        "paragraph": WD_STYLE_TYPE.PARAGRAPH,
        "character": WD_STYLE_TYPE.CHARACTER,
        "table": WD_STYLE_TYPE.TABLE,
        "list": WD_STYLE_TYPE.LIST
    }

    for style in doc.styles:
        # Filter by type if specified
        if style_type != "all":
            if style_type.lower() in type_map:
                if style.type != type_map[style_type.lower()]:
                    continue
            else:
                continue

        style_info = {
            "name": style.name,
            "type": str(style.type).replace("WD_STYLE_TYPE.", ""),
            "builtin": style.builtin
        }

        # Add base style if available
        if style.base_style:
            style_info["base_style"] = style.base_style.name

        styles_list.append(style_info)

    # Sort by type and name
    styles_list.sort(key=lambda x: (x["type"], x["name"]))

    return {
        "total_styles": len(styles_list),
        "filter": style_type,
        "styles": styles_list
    }


# ============================================
# Bookmark & Hyperlink Listing Tools
# ============================================

@mcp.tool()
async def list_bookmarks() -> dict:
    """
    List all bookmarks in the document.
    """
    doc = get_document()
    bookmarks = []

    # Find all bookmark start elements in the document
    for para_idx, para in enumerate(doc.paragraphs):
        bookmark_starts = para._element.findall('.//w:bookmarkStart', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        for bookmark in bookmark_starts:
            name = bookmark.get(qn('w:name'))
            bookmark_id = bookmark.get(qn('w:id'))
            if name and not name.startswith('_'):  # Skip internal bookmarks
                bookmarks.append({
                    "name": name,
                    "id": bookmark_id,
                    "paragraph_index": para_idx,
                    "paragraph_preview": para.text[:100] + "..." if len(para.text) > 100 else para.text
                })

    return {
        "total_bookmarks": len(bookmarks),
        "bookmarks": bookmarks
    }


@mcp.tool()
async def list_hyperlinks() -> dict:
    """
    List all hyperlinks in the document.
    """
    doc = get_document()
    hyperlinks = []

    for para_idx, para in enumerate(doc.paragraphs):
        # Find hyperlink elements
        hyperlink_elements = para._element.findall('.//w:hyperlink', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

        for hyperlink in hyperlink_elements:
            r_id = hyperlink.get(qn('r:id'))
            anchor = hyperlink.get(qn('w:anchor'))

            # Get the display text
            text_elements = hyperlink.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            display_text = ''.join(t.text for t in text_elements if t.text)

            # Try to get the URL from relationships
            url = None
            if r_id:
                try:
                    rel = para.part.rels[r_id]
                    url = rel.target_ref
                except:
                    pass

            hyperlinks.append({
                "display_text": display_text,
                "url": url or "(internal link)",
                "anchor": anchor,
                "paragraph_index": para_idx
            })

    return {
        "total_hyperlinks": len(hyperlinks),
        "hyperlinks": hyperlinks
    }


@mcp.tool()
async def remove_bookmark(bookmark_name: str) -> dict:
    """
    Remove a bookmark from the document.

    Args:
        bookmark_name: The name of the bookmark to remove
    """
    doc = get_document()
    removed = False

    # Find and remove bookmark start and end elements
    for para in doc.paragraphs:
        bookmark_starts = para._element.findall('.//w:bookmarkStart', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        for bookmark in bookmark_starts:
            name = bookmark.get(qn('w:name'))
            if name == bookmark_name:
                bookmark_id = bookmark.get(qn('w:id'))

                # Remove bookmark start
                bookmark.getparent().remove(bookmark)

                # Find and remove corresponding bookmark end
                for p in doc.paragraphs:
                    bookmark_ends = p._element.findall('.//w:bookmarkEnd', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    for end in bookmark_ends:
                        if end.get(qn('w:id')) == bookmark_id:
                            end.getparent().remove(end)

                removed = True
                break
        if removed:
            break

    if not removed:
        return {"error": f"Bookmark '{bookmark_name}' not found in the document."}

    save_document(doc)

    return {
        "status": "success",
        "message": f"Removed bookmark '{bookmark_name}'"
    }


@mcp.tool()
async def remove_hyperlink(query: str, threshold: float = 0.5) -> dict:
    """
    Remove a hyperlink from the document, keeping the display text.

    Args:
        query: Text to search for to find the paragraph containing the hyperlink
        threshold: Minimum similarity score 0-1 (default 0.5)
    """
    doc = get_document()

    match = find_paragraph_by_text(doc, query, threshold)
    if not match:
        return {"error": f"No paragraph found matching: '{query}'"}

    idx, para, score = match
    removed_count = 0

    # Find and remove hyperlinks in this paragraph
    hyperlinks = para._element.findall('.//w:hyperlink', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

    for hyperlink in hyperlinks:
        # Get all runs inside the hyperlink
        runs = hyperlink.findall('.//w:r', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

        # Move runs outside the hyperlink
        parent = hyperlink.getparent()
        index = list(parent).index(hyperlink)

        for run in runs:
            parent.insert(index, run)
            index += 1

        # Remove the empty hyperlink element
        parent.remove(hyperlink)
        removed_count += 1

    if removed_count == 0:
        return {"error": "No hyperlinks found in the matched paragraph."}

    save_document(doc)

    return {
        "status": "success",
        "message": f"Removed {removed_count} hyperlink(s) from paragraph {idx}",
        "paragraph_text": para.text[:100] + "..." if len(para.text) > 100 else para.text
    }


# ============================================
# Page Layout Tools
# ============================================

@mcp.tool()
async def set_page_margins(
    top: float = None,
    bottom: float = None,
    left: float = None,
    right: float = None,
    section_index: int = 0
) -> dict:
    """
    Set page margins for a document section.
    All measurements are in inches.

    Args:
        top: Top margin in inches (e.g., 1.0)
        bottom: Bottom margin in inches
        left: Left margin in inches
        right: Right margin in inches
        section_index: Index of the section to modify (default 0)
    """
    doc = get_document()

    if section_index >= len(doc.sections):
        return {"error": f"Section {section_index} not found. Document has {len(doc.sections)} section(s)."}

    section = doc.sections[section_index]
    updated = []

    if top is not None:
        section.top_margin = Inches(top)
        updated.append(f"top={top}\"")
    if bottom is not None:
        section.bottom_margin = Inches(bottom)
        updated.append(f"bottom={bottom}\"")
    if left is not None:
        section.left_margin = Inches(left)
        updated.append(f"left={left}\"")
    if right is not None:
        section.right_margin = Inches(right)
        updated.append(f"right={right}\"")

    if not updated:
        return {"error": "No margins specified to update."}

    save_document(doc)

    return {
        "status": "success",
        "message": f"Updated margins for section {section_index}: {', '.join(updated)}",
        "current_margins": {
            "top": section.top_margin.inches,
            "bottom": section.bottom_margin.inches,
            "left": section.left_margin.inches,
            "right": section.right_margin.inches
        }
    }


@mcp.tool()
async def set_page_size(
    width: float = None,
    height: float = None,
    orientation: str = None,
    preset: str = None,
    section_index: int = 0
) -> dict:
    """
    Set page size for a document section.

    Args:
        width: Page width in inches (ignored if preset is used)
        height: Page height in inches (ignored if preset is used)
        orientation: "portrait" or "landscape"
        preset: Preset page size - "letter", "a4", "legal", "a3", "a5"
        section_index: Index of the section to modify (default 0)
    """
    doc = get_document()

    if section_index >= len(doc.sections):
        return {"error": f"Section {section_index} not found. Document has {len(doc.sections)} section(s)."}

    section = doc.sections[section_index]

    # Page size presets (width x height in inches)
    presets = {
        "letter": (8.5, 11),
        "legal": (8.5, 14),
        "a4": (8.27, 11.69),
        "a3": (11.69, 16.54),
        "a5": (5.83, 8.27)
    }

    if preset:
        preset_lower = preset.lower()
        if preset_lower not in presets:
            return {"error": f"Unknown preset '{preset}'. Available: {', '.join(presets.keys())}"}
        width, height = presets[preset_lower]

    if width is not None:
        section.page_width = Inches(width)
    if height is not None:
        section.page_height = Inches(height)

    if orientation:
        if orientation.lower() == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            # Swap dimensions if needed
            if section.page_width < section.page_height:
                section.page_width, section.page_height = section.page_height, section.page_width
        elif orientation.lower() == "portrait":
            section.orientation = WD_ORIENT.PORTRAIT
            # Swap dimensions if needed
            if section.page_width > section.page_height:
                section.page_width, section.page_height = section.page_height, section.page_width

    save_document(doc)

    return {
        "status": "success",
        "message": f"Updated page size for section {section_index}",
        "current_size": {
            "width": round(section.page_width.inches, 2),
            "height": round(section.page_height.inches, 2),
            "orientation": "landscape" if section.orientation == WD_ORIENT.LANDSCAPE else "portrait"
        }
    }


@mcp.tool()
async def set_paragraph_spacing(
    query: str,
    line_spacing: float = None,
    space_before: float = None,
    space_after: float = None,
    threshold: float = 0.5
) -> dict:
    """
    Set spacing for a paragraph.

    Args:
        query: Text to search for to find the paragraph
        line_spacing: Line spacing multiplier (e.g., 1.0 for single, 1.5, 2.0 for double)
        space_before: Space before paragraph in points (e.g., 12)
        space_after: Space after paragraph in points (e.g., 12)
        threshold: Minimum similarity score 0-1 (default 0.5)
    """
    doc = get_document()

    match = find_paragraph_by_text(doc, query, threshold)
    if not match:
        return {"error": f"No paragraph found matching: '{query}'"}

    idx, para, score = match
    updated = []

    if line_spacing is not None:
        para.paragraph_format.line_spacing = line_spacing
        updated.append(f"line_spacing={line_spacing}")

    if space_before is not None:
        para.paragraph_format.space_before = Pt(space_before)
        updated.append(f"space_before={space_before}pt")

    if space_after is not None:
        para.paragraph_format.space_after = Pt(space_after)
        updated.append(f"space_after={space_after}pt")

    if not updated:
        return {"error": "No spacing values specified to update."}

    save_document(doc)

    return {
        "status": "success",
        "message": f"Updated spacing for paragraph {idx}: {', '.join(updated)}",
        "paragraph_text": para.text[:100] + "..." if len(para.text) > 100 else para.text
    }


# ============================================
# Section Tools
# ============================================

@mcp.tool()
async def get_sections() -> dict:
    """
    List all sections in the document with their properties.
    """
    doc = get_document()
    sections_list = []

    for idx, section in enumerate(doc.sections):
        section_info = {
            "index": idx,
            "page_width": round(section.page_width.inches, 2),
            "page_height": round(section.page_height.inches, 2),
            "orientation": "landscape" if section.orientation == WD_ORIENT.LANDSCAPE else "portrait",
            "margins": {
                "top": round(section.top_margin.inches, 2),
                "bottom": round(section.bottom_margin.inches, 2),
                "left": round(section.left_margin.inches, 2),
                "right": round(section.right_margin.inches, 2)
            },
            "has_header": bool(section.header.paragraphs and any(p.text for p in section.header.paragraphs)),
            "has_footer": bool(section.footer.paragraphs and any(p.text for p in section.footer.paragraphs))
        }
        sections_list.append(section_info)

    return {
        "total_sections": len(sections_list),
        "sections": sections_list
    }


@mcp.tool()
async def add_section_break(
    query: str = None,
    break_type: str = "next_page",
    threshold: float = 0.5
) -> dict:
    """
    Add a section break to the document.

    Args:
        query: Optional text to search for - adds section break after matching paragraph.
               If not provided, adds at the end of the document.
        break_type: Type of section break - "next_page", "continuous", "even_page", "odd_page"
        threshold: Minimum similarity score 0-1 (default 0.5)
    """
    doc = get_document()

    break_types = {
        "next_page": WD_SECTION.NEW_PAGE,
        "continuous": WD_SECTION.CONTINUOUS,
        "even_page": WD_SECTION.EVEN_PAGE,
        "odd_page": WD_SECTION.ODD_PAGE
    }

    if break_type.lower() not in break_types:
        return {"error": f"Unknown break type '{break_type}'. Available: {', '.join(break_types.keys())}"}

    if query:
        match = find_paragraph_by_text(doc, query, threshold)
        if not match:
            return {"error": f"No paragraph found matching: '{query}'"}
        idx, para, score = match

        # Add a paragraph after the matched one
        new_para = doc.add_paragraph()
        para._element.addnext(new_para._element)

        # Get the section and set the break type
        # We need to add section properties to the paragraph
        sectPr = OxmlElement('w:sectPr')
        type_elem = OxmlElement('w:type')
        type_elem.set(qn('w:val'), break_type.lower().replace('_', ''))
        sectPr.append(type_elem)
        new_para._element.append(sectPr)

        message = f"Section break ({break_type}) added after paragraph {idx}"
    else:
        # Add section break at end
        new_section = doc.add_section(break_types[break_type.lower()])
        message = f"Section break ({break_type}) added at end of document"

    save_document(doc)

    return {
        "status": "success",
        "message": message,
        "total_sections": len(doc.sections)
    }


@mcp.tool()
async def insert_line_break(query: str, after_text: str = None, threshold: float = 0.5) -> dict:
    """
    Insert a soft line break (Shift+Enter) within a paragraph.
    This creates a new line without starting a new paragraph.

    Args:
        query: Text to search for to find the paragraph
        after_text: Optional text after which to insert the line break.
                    If not provided, adds at the end of the paragraph.
        threshold: Minimum similarity score 0-1 (default 0.5)
    """
    doc = get_document()

    match = find_paragraph_by_text(doc, query, threshold)
    if not match:
        return {"error": f"No paragraph found matching: '{query}'"}

    idx, para, score = match

    if after_text:
        # Find the run containing the text and insert break after it
        found = False
        for run in para.runs:
            if after_text in run.text:
                # Split the run at the specified text
                parts = run.text.split(after_text, 1)
                if len(parts) == 2:
                    run.text = parts[0] + after_text
                    # Add line break
                    run.add_break(WD_BREAK.LINE)
                    # Add remaining text in a new run
                    if parts[1]:
                        new_run = para.add_run(parts[1])
                found = True
                break

        if not found:
            return {"error": f"Text '{after_text}' not found in the paragraph."}
    else:
        # Add line break at the end of the last run
        if para.runs:
            para.runs[-1].add_break(WD_BREAK.LINE)
        else:
            run = para.add_run()
            run.add_break(WD_BREAK.LINE)

    save_document(doc)

    return {
        "status": "success",
        "message": f"Line break inserted in paragraph {idx}",
        "paragraph_text": para.text[:100] + "..." if len(para.text) > 100 else para.text
    }


# ============================================
# Image Management Tools
# ============================================

@mcp.tool()
async def delete_image(query: str, image_index: int = 0, threshold: float = 0.5) -> dict:
    """
    Delete an image from the document.

    Args:
        query: Text to search for to find the paragraph containing the image
        image_index: If multiple images in paragraph, which one to delete (0-indexed, default 0)
        threshold: Minimum similarity score 0-1 (default 0.5)
    """
    doc = get_document()

    match = find_paragraph_by_text(doc, query, threshold)
    if not match:
        return {"error": f"No paragraph found matching: '{query}'"}

    idx, para, score = match

    # Find images (drawings) in the paragraph
    images = []
    for run in para.runs:
        drawings = run._element.findall('.//a:blip', {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
        if drawings:
            images.append((run, drawings))

    if not images:
        return {"error": "No images found in the matched paragraph."}

    if image_index >= len(images):
        return {"error": f"Image index {image_index} out of range. Paragraph has {len(images)} image(s)."}

    # Delete the run containing the image
    run_to_delete, _ = images[image_index]
    run_to_delete._element.getparent().remove(run_to_delete._element)

    save_document(doc)

    return {
        "status": "success",
        "message": f"Deleted image {image_index} from paragraph {idx}",
        "remaining_images": len(images) - 1
    }


if __name__ == "__main__":
    # Ensure the documents directory exists
    os.makedirs("documents", exist_ok=True)

    # Ensure the default document exists or create a blank one
    if not os.path.exists(CURRENT_DOCX_PATH):
        print(f"Creating new document at {CURRENT_DOCX_PATH}")
        doc = Document()
        doc.save(CURRENT_DOCX_PATH)

    print("=" * 60)
    print("        DOCX MCP SERVER")
    print("=" * 60)
    print(f"\nCurrent document: {os.path.abspath(CURRENT_DOCX_PATH)}")
    print(f"\n🌐 Server: http://localhost:8000")
    print(f"   SSE Endpoint: http://localhost:8000/sse")

    print("\n📝 Quick Tips:")
    print("   - Use 'switch_document' to work on different files")
    print("   - 60+ tools available for document editing")
    print("=" * 60)
    print()

    # Run the server
    mcp.run(
        transport="sse",
        host="0.0.0.0",
        port=8000
    )
